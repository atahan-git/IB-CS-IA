# -*- coding: utf-8 -*-
#coding:utf8

import datetime
import threading;
from tkinter.filedialog import *
from tkinter.messagebox import *

import xlrd;
from docx import Document
from xlrd import open_workbook
from xlutils.copy import copy

# -------------------------------File Names
if getattr (sys, 'frozen', False):
    # frozen
    myPath = os.path.dirname (sys.executable)
else:
    # unfrozen
    myPath = os.path.dirname (os.path.realpath (__file__))
dictionary = "Dictionary.xls";
dictionaryTemp = "Dictionary_temp.xls"
inputFixed = "input_Fixed.xls"
input_def = "file";
outputTemplate = "Output Template.docx";
output = "Output.docx";

# -------------------------------Dictionary Constants
dictionary_normal_x = 3;
dictionary_normal_y = 1 + 8;

dictionary_specific_x = 3;
dictionary_specific_y = 4 + 8;

endOffTheLineValue = "__end_of_the_line__";
endOffTheColumnValue = "__switch_columns__";
dynamicFill = "__dynamic_fill_slot__";
split = "__";
averageGrade = "avg";

dictionary_columnChangeOffset = 4;

# -------------------------------Input.xls constants
inputXlsYOffset = 9;
inputXlsXGradeSearchBegin = 24;
inputXlsXOffset = 6;


# -------------------------------File Set up Related Methods
def OpenBaseDocuments ():
    global wb_dic, dic, wb_dic_write, dic_write, doc, tb_names, tb_IB, tb;
    try:
        wb_dic = open_workbook (dictionary);
        dic = wb_dic.sheets ()[0]  ## possible error throw
        wb_dic_write = copy (wb_dic);
        dic_write = wb_dic_write.get_sheet (0)  ## possible error throw

        doc = Document (outputTemplate);  ## possible error throw
        tb_names = doc.tables[0]
        tb_IB = doc.tables[1]
        tb = doc.tables[2]
    except:
        return FileReadorWriteError (OpenBaseDocuments);


OpenBaseDocuments();


def FileReadorWriteError (callback):  # this is a critical error, we try to do it again and exit if canceled
    print()
    print(str(callback))
    if (askretrycancel ("I/O Error", "There was an error while reading or writing files. " +
                                     "Please try closing any open files, checking your antivirus, running with administrator rights, or reinstalling the program to fix this issue.")):
        return callback ();
    else:
        root.destroy ();
        sys.exit ();
        return False;


def SetupInput ():
    global wb, ws, isInputProper;
    try:
        wb = open_workbook (input);
        ws = wb.sheets ()[0]  ## possible error throw
        return True;
    except:
        showwarning (title = "Unrecognized File Format", message = "Please select a proper E-Okul Excel file.")
        inputDisp.set (input_def);
        isInputProper = False;
        return False;


def SelectFile ():
    global input;
    global isInputProper;
    input = askopenfilename (initialdir = "myPath", title = "Please Select E-okul Excel File");
    if(not(SetupInput ())):
        return;
    try:
        # if there isn't this specific text in this specific cell, this is the wrong file. Show an error
        if (ws.cell_value (inputXlsXOffset, 0) != "ÖĞRENCİNİN"):
            showwarning (title = "Wrong E-Okul file", message = "Please select the correct E-Okul Excel file. You probably got the 'data only' Excel file. You need the other one.")
            inputDisp.set (input_def);
            isInputProper = False;
            return;
    except:  # if we can't read the file show an error
        showwarning (title = "Unrecognized File Format", message = "Please select a proper E-Okul Excel file.");
        inputDisp.set (input_def);
        isInputProper = False;
        return;
    try:
        # if this slot has a number less than 10 then this slot is a "hours" slot. This means that the lesson grades are shifted, which needs fixing
        if (float (ws.cell_value (inputXlsXGradeSearchBegin+2, 8)) < 10):
            FixInput ();
    except Exception as e:
        print (e)
        pass;
    # print(input);
    inputDisp.set (input[input.rindex ("/") + 1:]);

    print ("The input is proper!")
    isInputProper = True;


fix_x_start = inputXlsXGradeSearchBegin;
fix_x_end = 91;
fix_y_start = 8;
fix_y_end = 16;


def FixInput ():
    global wb, ws, wb_write;
    print ("Trying to fix")
    try:
        wb_write = copy (wb);
        ws_write = wb_write.get_sheet (0)  ## possible error throw
    except:
        return FileReadorWriteError (FixInput);

    for x in range (fix_x_start, fix_x_end, 1):
        for y in range (fix_y_start, fix_y_end + 1, 2):
            try:
                # print("MY Y Checks " + str(x) + " - " +  str(y) + " - " + str(ws.cell_value (x, y)));
                if (float (ws.cell_value (x, y)) < 10):
                    ShiftCellsLeft (ws_write, x, y)
            except Exception as e:
                if (ws.cell_value (x, y) != ""):
                    print ("shift cells exception")
                    print (e)
                pass;

    SaveAndReOpenFile();

    if (askyesno ("Data Validation", "Some of the lessons didn't match correctly with their hours/grades grid in the excel file. Automatic corrections were made. Try running the program one time to see if it works, and if there are problems check the file before running it again. Open the corrected file now?")):
        to = threading.Thread (target = OpenFixedFile)
        to.start ()


def SaveAndReOpenFile ():
    global wb, ws, wb_write;
    try:
        wb_write.save (inputFixed);
        wb = open_workbook (inputFixed);
        ws = wb.sheets ()[0]  ## possible error throw
    except:
        return FileReadorWriteError (SaveAndReOpenFile);



def OpenFixedFile ():
    os.startfile (inputFixed)


def ShiftCellsLeft (ws_write, row, column):
    global wb;
    global ws;
    print ("Shifting to " + str (column) + " row:" + str (row))
    for y in range (fix_y_end, column, -1):
        try:
            print ("Shifted " + str (row) + " - " + str (y) + " == " + str (ws.cell_value (row, y - 1)) + " > " + str (ws.cell_value (row, y)))
            ws_write.write (row, y, str (ws.cell_value (row, y - 1)));
        except Exception as e:
            print (e)
            pass;
    ws_write.write (row, column, "");
    SaveAndReOpenFile();

# -------------------------------Helper Methods
def FindRowInXLS (item):  # finds the item in our input.xls
    rowX = inputXlsXGradeSearchBegin;
    isFound = False;

    try:
        while not (isFound):
            if (ws.cell_value (rowX, 0) == item):
                isFound = True;
            else:
                rowX += 1;
    except:
        return -1;

    return rowX;


def FindRowInDOCX (item):  # finds the item in our output.docx
    rowX = 0;
    isFound = False;

    try:
        while not (isFound):
            # print(item + " - " + tb.cell (rowX, 0).text);
            if (tb.cell (rowX, 0).text == item):
                isFound = True;
            else:
                rowX += 1;
    except:
        return -1;

    return rowX;


def WriteInDOCX (text, x, y):  # writes text into the given row and column without editing text formatting
    try:
        paragraph = tb.cell (x, y).paragraphs[0]
        paragraph.runs[0].text = str (text);
    except:
        print ("-" * 50 + "ILLEGAL WRITE REQUEST --> " + str (text));


def RemoveDOCXRow (table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove (tr)


def WriteInDOCX_IB (text, x, y):  # writes text into the given row and column without editing text formatting - for IB lessons
    try:
        paragraph = tb_IB.cell (x, y).paragraphs[0]
        paragraph.runs[0].text = str (text);
    except:
        print ("-" * 50 + "ILLEGAL WRITE REQUEST IB--> " + str (text));


def WriteInDOCX_Names (text, x, y):  # writes text into the given row and column without editing text formatting - for name and other stuff
    try:
        paragraph = tb_names.cell (x, y).paragraphs[0]
        paragraph.runs[0].text = str (text);
    except:
        print ("-" * 50 + "ILLEGAL WRITE REQUEST Names--> " + str (text));


def WriteInDictionary (text, x, y):  # writes text into the given row and column without editing text formatting
    try:
        dic_write.write (x, y, text);
        print("Written in dictionary: " + text + " - " + str(x) + "," + str(y))
    except:
        print ("-" * 25 + "ILLEGAL DICTIONARY WRITE --> " + str (text));


def InitializeSelectionWithSettings (label, frame, options, packX,
                                     gridSize):  # initializes a input field with options to choose from and returns the variable
    tempVar = StringVar (root);
    tempVar.set (defaultOption);  # default value
    options = OptionMenu (frame, tempVar, *options);
    options.config (width = gridSize);
    Label (frame, text = label).grid (row = packX, column = 0, sticky = (EW), ipadx = 5, ipady = 5);
    options.grid (row = packX, column = 1, sticky = (EW), ipadx = 5);

    return tempVar;


# -------------------------------GUI and variable set up
defaultOption = "...";

root = Tk ()
root.title ("Converter 1.1.2 12th Grade Edition")
root.resizable (width = False, height = False)
root.iconbitmap ('icon.ico')

# Add a grid
mainframe = Frame (root)
mainframe.config (width = 500);
mainframe.grid (column = 0, row = 0, sticky = (N, W, E, S))
mainframe.columnconfigure (0, weight = 1)
mainframe.rowconfigure (0, weight = 1)
mainframe.pack (pady = 50, padx = 20)

Label (mainframe, text = "E-Okul to IB Transcript Converter", font = ("TkDefaultFont", 15, "bold")).grid (row = 0, column = 0,sticky = (EW), ipadx = 20);
Label (mainframe).grid (row = 1, column = 0, sticky = (EW), ipady = 2);  # spacer

isInputProper = False;

input = "input.xls";
inputDisp = StringVar (root);
inputDisp.set (input_def);

fileFrame = Frame (mainframe);
fileFrame.grid (row = 2, column = 0);
Label (fileFrame, text = "Select file: ").grid (row = 0, column = 0, ipadx = 10);
Label (fileFrame, textvariable = inputDisp, relief = GROOVE, width = 25).grid (row = 0, column = 1, ipadx = 10);
Label (fileFrame, text = "").grid (row = 0, column = 2);  # spacer
Button (fileFrame, text = "Browse", command = SelectFile).grid (row = 0, column = 3, ipadx = 20);

normalLessons = Frame (mainframe)
normalLessons.grid (row = 3, column = 0);

foreignLang = InitializeSelectionWithSettings ("2nd Foreign Language: ", normalLessons, ["FRENCH", "GERMAN"], 0, 20);

ibLessons = [];

Label (mainframe).grid (row = 4, column = 0, sticky = (EW), ipady = 10);  # spacer
Label (mainframe, text = "IB Lessons", font = ("TkDefaultFont", 13)).grid (row = 5, column = 0, sticky = (EW), ipadx = 20);
Label (mainframe).grid (row = 6, column = 0, sticky = (EW), ipady = 0);  # spacer

ibFrame = Frame (mainframe)
ibFrame.grid (row = 7, column = 0);

i = 0;  # lesson id
in_x = -1;  # keep track of the row we read
# start from IB Column
y_track = dictionary_columnChangeOffset;

firstSpace = False;  # first space designates a new lesson
endOfTheLine = False;  # second space designates end of the lessons
while not (endOfTheLine):  # repeat until there are no more lessons
    read_name = "";
    firstSpace = False;
    thisLessonOptions = []
    # print ("Reading IB Lessons")
    while not (firstSpace):  # ----------------------------------------------------------------READ LESSON NAMES FROM XLS FILE
        in_x += 1;
        read_name = dic.cell_value (dictionary_normal_x + in_x, dictionary_normal_y + y_track + 1);
        # print (read_name + " *");
        if(read_name == ""):
            continue;
        if (read_name == split or read_name == dynamicFill):
            firstSpace = True;
        else:
            thisLessonOptions.append (read_name);

    if (len (thisLessonOptions) > 0):
        print (thisLessonOptions)
        ibLessons.append (InitializeSelectionWithSettings ("Group " + str (i + 1) + ":", ibFrame, thisLessonOptions, i, 40));
        i += 1;

    if (dic.cell_value (dictionary_normal_x + in_x + 1, dictionary_normal_y + y_track) == endOffTheLineValue):
        endOfTheLine = True;
        print ("IB Lesson Options Read Complete")
        print ();
        break;

convertText = "Converting...  ";
buttonText = StringVar (root);
buttonText.set (convertText)


def ok ():
    t = threading.Thread (target = SetUpDictionary)
    t.start ()
    t2 = threading.Thread (target = WriteStudentInfo)
    t2.start ()


# ------------------------------- Write name & surname & ID etc.
def WriteStudentInfo ():  # this wont change for anyone or throughout the years (unless they change the whole file) so there is no point in writing these in a dictionary

    WriteInDOCX_Names (ws.cell_value (inputXlsXOffset + 1, 3), 0, 1);
    WriteInDOCX_Names (ws.cell_value (inputXlsXOffset + 2, 3), 0, 3);
    WriteInDOCX_Names (ws.cell_value (inputXlsXOffset + 3, 3), 1, 1);
    WriteInDOCX_Names (ws.cell_value (inputXlsXOffset + 4, 3), 2, 1);
    WriteInDOCX_Names (ws.cell_value (inputXlsXOffset + 5, 3), 1, 5);
    WriteInDOCX_Names (ws.cell_value (inputXlsXOffset + 2, 10), 0, 5);
    WriteInDOCX_Names (ConvertDate (datetime.datetime (*xlrd.xldate_as_tuple (ws.cell_value (inputXlsXOffset + 3, 10), wb.datemode)).date ().isoformat ()), 1, 3);
    WriteInDOCX_Names ("Male" if ws.cell_value (inputXlsXOffset + 4, 10) == "Erkek" else "Female", 2, 3);
    gradevalue = "-MANUAL ENTRY-";
    myVal = ws.cell_value (inputXlsXOffset + 5, 10);
    if myVal == "AL - 11. Sınıf / A Şubesi":
        gradevalue = "11th Grade /A Branch"
    elif myVal == "AL - 11. Sınıf / B Şubesi":
        gradevalue = "11th Grade /B Branch"
    elif myVal == "AL - 12. Sınıf / A Şubesi":
        gradevalue = "12th Grade /A Branch"
    elif myVal == "AL - 12. Sınıf / B Şubesi":
        gradevalue = "12th Grade /B Branch"
    WriteInDOCX_Names (gradevalue, 2, 5);


def ConvertDate (date):
    parts = date.split ("-");

    return parts[2] + "/" + parts[1] + "/" + parts[0];

physics = "SEÇMELİ FİZİK";
chem = "SEÇMELİ KİMYA";
bio = "SEÇMELİ BİYOLOJİ";
def FixScienceLesson (xOffset, yOffset, scienceName):
    print("Fixing science lesson: " + str(xOffset) + " - " + str(yOffset) + " - " + scienceName)
    isFirst = True;
    yOffset -= 1;
    x_in = 0;
    while(dic.cell_value(xOffset + x_in, yOffset) != split):
        searchName = dic.cell_value(xOffset + x_in,yOffset);
        print(str(xOffset + x_in) + " - " + str(yOffset) + " - " + str(searchName))
        if not(searchName == physics or searchName == chem or searchName == bio):
            x_in -= 1;
        else:
            WriteInDictionary("",xOffset + x_in, yOffset);
            if(isFirst):
                WriteInDictionary (scienceName, xOffset + x_in, yOffset);
                isFirst = False;

            x_in -= 1;


world = "SEÇMELİ ÇAĞDAŞ DÜNYA TARİHİ";
def FixDoubleWorldHistory (xOffset, yOffset, n):
    print("Fixing double world history: " + str(xOffset) + " - " + str(yOffset) + " - " + str(n))
    yOffset -= 1;
    x_in = 0;
    isFirst = True;
    while (True):
        print(str(xOffset + x_in) + " - " + str(yOffset))
        searchName = dic.cell_value (xOffset + x_in, yOffset);
        print(searchName)
        if (searchName != world):
            if n < 4:
                x_in += 1;
            else:
                x_in -= 1;

        else:
            if not (n>4 and isFirst):
                WriteInDictionary ("", xOffset + x_in, yOffset);
                return;
            else:
                x_in -= 1;
                isFirst = False;


# ------------------------------- Set up the dictionary file according to our lesson selection
def SetUpDictionary ():
    global dic;
    if (not isInputProper):
        showerror ("Error", "Please select a proper input file!")
        print ("input not set!")
        return;
    if (foreignLang.get () == defaultOption):
        showerror ("Error", "Please select all the lessons!")
        print ("not all vars selected!")
        return;
    for stvar in ibLessons:
        if (stvar.get () == defaultOption):
            showerror ("Error", "Please select all the lessons!")
            print ("not all vars selected!")
            return;
    button.config (state = DISABLED, textvariable = buttonText, disabledforeground = "GREY")  # configure our button to be a makeshift progress bar

    # change our dictionary file according to the settings choosen
    WriteInDictionary (foreignLang.get (), dictionary_normal_x, dictionary_normal_y + 1);
    WriteInDictionary ("ELECTIVE FRENCH" if foreignLang.get () == "FRENCH" else "ELECTIVE GERMAN", dictionary_normal_x + 2, dictionary_normal_y + 1);

    ##set every ib lesson according to users selection both in our dictionary and word file
    in_x = -1;
    n = -1;
    for lesson in ibLessons:
        in_x += 1;
        n += 1;

        WriteInDOCX_IB (lesson.get (), n + 2, 0);

        while (True):
            print(dic.cell_value (dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1))
            if (dic.cell_value (dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1) == dynamicFill):
                WriteInDictionary (lesson.get (), dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1);
                print ("written value " + lesson.get ())
                #--------------------------------------------------------------------------------Science lessons fix for our double science students
                if(lesson.get () == "PHYSICS SL" or lesson.get () == "PHYSICS HL"):
                    FixScienceLesson(dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1, physics);
                elif (lesson.get () == "CHEMISTRY SL" or lesson.get () == "CHEMISTRY HL"):
                    FixScienceLesson (dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1, chem);
                elif (lesson.get () == "BIOLOGY SL" or lesson.get () == "BIOLOGY HL"):
                    FixScienceLesson (dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1, bio);
                elif(n == 5):
                    FixScienceLesson(dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1, "");
                if (lesson.get () == "WORLD HISTORY SL" or lesson.get () == "WORLD HISTORY HL"):
                    FixDoubleWorldHistory (dictionary_normal_x + in_x, dictionary_normal_y + dictionary_columnChangeOffset + 1, n);
                break;
            else:
                in_x += 1;

    try:
        wb_dic_write.save (dictionaryTemp);
        wb_dic = open_workbook (dictionaryTemp);
        dic = wb_dic.sheets ()[0]  ## possible error throw
    except:
        return FileReadorWriteError (SetUpDictionary);

    # setting up dictionary file complete
    print ();
    print ("-*-*-*-*-*-*-*-*-*-*- Dictionary Set-up Complete! -*-*-*-*-*-*-*-*-*-*-");
    print ();

    ConvertLessons ();
##------------------



##-------

# -------------------------------Main Method
def ConvertLessons ():
    for g in range (0, 4):  # repeat all this for 9 through 12 grade
        #if (g == 3):  # SKIP 12TH GRADE BECAUSE IT DOESNT APPLY TO US
        #    continue;

        global buttonText;
        buttonText.set (convertText + str (g * 20) + "%");

        print ("-----------------------------------------------------------")
        print ("------------------ Starting for grade: " + str (9 + g) + " -----------------");
        print ("-----------------------------------------------------------")
        read_y = inputXlsYOffset + g * 2;  # get correct field y value - every grade has 2 fields so i*2

        in_x = -1;  # keep track of the row we read
        cur_ibLesson = 0;
        # our dictionary file has 2 different y offsets for different columns,
        # this variable keeps track of which track we are currently in
        y_track = 0;

        firstSpace = False;  # first space designates a new lesson
        endOfTheLine = False;  # second space designates end of the lessons
        while not (endOfTheLine):  # repeat until there are no more lessons
            t_hours = 0;  # total hours this lesson have
            t_grade = 0;  # total grade this lesson have - we need to average it later
            read_name = "";
            read_rowX = 0;
            firstSpace = False;
            print ("**-- ==> Checking Values --**")
            while not (firstSpace):  # ----------------------------------------------------------------READ GRADES FROM XLS FILE
                in_x += 1;
                try:
                    read_name = dic.cell_value (dictionary_normal_x + in_x, dictionary_normal_y + y_track);
                except:
                    print ("read failure" + str (dictionary_normal_x + in_x) + " - " + str (dictionary_normal_y + y_track));

                if(read_name == ""):
                    continue;
                if (read_name == split):
                    firstSpace = True;
                else:
                    print (read_name + " *");
                    read_rowX = FindRowInXLS (read_name);
                    if (read_rowX != -1):  # check if the row actually exists
                        try:  # check if there is an actual value for this lesson - if not pass this lesson.
                            t_hours += int (float (ws.cell_value (read_rowX, read_y)));
                            t_grade += int (float (ws.cell_value (read_rowX, read_y))) * float (ws.cell_value (read_rowX, read_y + 1));
                        except Exception as e:
                            print ("Error: " + str(ws.cell_value (read_rowX, read_y)) + " - " + str(ws.cell_value (read_rowX, read_y + 1)));
                            print(e)
                            pass;

            # ---------------------------------------------------------------------------------------WRITE GRADES TO DOCX FILE
            # Note: IB lessons won't be written
            # NON IB LESSONS
            if (y_track == 0):
                write_name = dic.cell_value (dictionary_normal_x + in_x - 1, dictionary_normal_y + 1);

                print ("--** Writing Values ==> **--")
                if (t_hours > 0):  # no need to write lessons that have no class time
                    print (str (g) + "-" + str (in_x) + " -> " + write_name + " ==> " + str (t_hours) + " - " + str (
                        '%.2f' % (t_grade / t_hours)));
                    write_rowX = FindRowInDOCX (write_name);
                    WriteInDOCX (t_hours, write_rowX, g * 2 + 1);
                    WriteInDOCX ('%.2f' % (t_grade / t_hours), write_rowX, g * 2 + 2);
                else:
                    if (write_name != ""):
                        print (str (in_x) + " -> " + write_name + " ==> " + "0 hours detected");

            # IB LESSONS
            elif (g > 1):
                write_name = dic.cell_value (dictionary_normal_x + in_x - 1, dictionary_normal_y + y_track + 1);

                print ("--** Writing Values ==> **--")
                if (t_hours > 0):  # no need to write lessons that have no class time
                    print ("IB" +"-"+str (in_x) + " -> " + write_name + " ==> " + str (t_hours) + " - " + str (
                        '%.2f' % (t_grade / t_hours)));
                    WriteInDOCX_IB (t_hours, cur_ibLesson + 2, (g - 2) * 2 + 1);
                    WriteInDOCX_IB ('%.2f' % (t_grade / t_hours), cur_ibLesson + 2, (g - 2) * 2 + 2);
                    cur_ibLesson += 1;
                else:
                    if (write_name != ""):
                        print (str (in_x) + " -> " + write_name + " ==> " + "0 hours detected");

            # if we get the end of the line element then we got to the end of the list
            if (dic.cell_value (dictionary_normal_x + in_x + 1, dictionary_normal_y + y_track) == endOffTheLineValue):
                endOfTheLine = True;
                print ("Stopping checking for this grade")
                break;

            if (dic.cell_value (dictionary_normal_x + in_x + 1, dictionary_normal_y + y_track) == endOffTheColumnValue):
                if (g > 1):  # IB Lessons only exists in 11&12th grade - don't change to IB track if we are not in those grades
                    y_track = dictionary_columnChangeOffset;
                    in_x = -1;
                    print ("Switched tracks")
                else:
                    endOfTheLine = True;
                    print ("Stopping checking for this grade")
                    break;
            print ();
    # --------------------------------------------------------------------------------------- CLEANUP EMPTY LESSONS
    print ()
    print ("Cleaning up empty lessons")
    x = 2;
    x -= 1;
    while (tb.cell (x, 0).text != "Total Number of Weekly Class Hours"):
        skipRow = False;
        x += 1;
        print (x)
        print (tb.cell (x, 0).text)
        for y in range (1, 8):
            myText = tb.cell (x, y).text
            if (myText != "0" and myText != "" and myText != "2"):
                #x -= 1;
                skipRow = True;
                continue;
        if(not skipRow):
            print ("Removing row: " + str (tb.cell (x, 0).text))
            RemoveDOCXRow (tb, tb.rows[x]);
            x -= 1;

    buttonText.set (convertText + str (80) + "%");
    # --------------------------------------------------------------------------------------- CALCULATE AVERAGE GRADES & LESSON HOURS
    print ()
    print ("Calculating averages")
    # NON IB LESSONS
    for g in range (0, 4):
        if (g == 3):  # SKIP 12TH GRADE BECAUSE IT DOESNT APPLY TO US
            continue;

        global avgX;
        t_hours = 0;
        t_grade = 0;
        x = 2;
        while (tb.cell (x, g * 2 + 1).text != averageGrade):  # add up grades and hours of all the lesson until avg slot
            # print (str (x) + " ==> " + str (tb.cell (x, g * 2).text) + " - " + str (tb.cell (x, g * 2 + 1).text) + " - " + str (tb.cell (x, g * 2 + 2).text))
            if (tb.cell (x, g * 2 + 2).text != ""):
                t_hours += int (tb.cell (x, g * 2 + 1).text);
                t_grade += int (tb.cell (x, g * 2 + 1).text) * float (tb.cell (x, g * 2 + 2).text);
            x += 1;
        avgX = x;
        if (t_hours > 0):  # write found values to the file
            if (g == 0 or g == 1 or g == 2):  # add the guidance hour if there is a guidance lesson at that grade
                print (str (9 + g) + " grade total hours: " + str (t_hours + 1) + " average grade: " + str (
                    '%.2f' % (t_grade / t_hours)));
                WriteInDOCX (t_hours + 1, x, g * 2 + 1);
            else:
                print (str (9 + g) + " grade total hours: " + str (t_hours) + " average grade: " + str (
                    '%.2f' % (t_grade / t_hours)));
                WriteInDOCX (t_hours, x, g * 2 + 1);
            WriteInDOCX ('%.2f' % (t_grade / t_hours), x + 1, g * 2 + 2);
        else:
            print ("ERROR - there isnt any grades for this grade: " + str (9 + g));

    print ()
    print ("Calculating IB averages")
    # IB LESSONS
    for g in range (0, 2):
        if(g == 1): #SKIP 12TH GRADE BECAUSE IT DOESNT APPLY TO US
            continue;

        t_hours = 0;
        t_grade = 0;
        tokAdd = 0;
        x = 2;
        while (tb_IB.cell (x, g * 2 + 1).text != averageGrade):  # add up grades and hours of all the lesson until avg slot
            print (str (x) + " ==> " + str (tb_IB.cell (x, 0).text) + " - " + str (tb_IB.cell (x, g * 2 + 1).text) + " - " + str (tb_IB.cell (x, g * 2 + 2).text))
            if (tb_IB.cell (x, g * 2 + 2).text != "N/A"):
                if (tb_IB.cell (x, g * 2 + 2).text != ""):
                    t_hours += int (tb_IB.cell (x, g * 2 + 1).text);
                    t_grade += int (tb_IB.cell (x, g * 2 + 1).text) * float (tb_IB.cell (x, g * 2 + 2).text);
            else:
                print ("TOK Added")
                if(g == 0):
                    tokAdd = 1;
                else:
                    tokAdd = 2;
            x += 1;

        if (t_hours > 0):  # write found values to the file
            print (str (11 + g) + " grade total IB hours: " + str (t_hours + tokAdd) + " average IB grade: " + str (
                '%.2f' % (t_grade / t_hours)));
            WriteInDOCX_IB (t_hours + tokAdd, x, g * 2 + 1);
            WriteInDOCX_IB ('%.2f' % (t_grade / t_hours), x, g * 2 + 2);
        else:
            print ("ERROR - there isnt any grades for this IB grade: " + str (9 + g));

    print ()
    print ("Calculating General Averages")
    avgX += 1;

    weightedGPA = [];
    weightedGPA.append (((int (tb.cell (avgX - 1, 5).text) - 1) * float (tb.cell (avgX, 6).text) + (int (tb_IB.cell (9, 1).text) - 1) * float (tb_IB.cell (9, 2).text)) / (
            (int (tb.cell (avgX - 1, 5).text) - 1) + (int (tb_IB.cell (9, 1).text)) - 1));
    try:
        weightedGPA.append (((int (tb.cell (avgX - 1, 7).text) - 1) * float (tb.cell (avgX, 8).text) + (int (tb_IB.cell (9, 3).text) - 2) * float (tb_IB.cell (9, 4).text)) / (
                (int (tb.cell (avgX - 1, 7).text) - 1) + (int (tb_IB.cell (9, 3).text)) - 2));
    except:
        weightedGPA.append(0);

    WriteInDOCX ('%.2f' % weightedGPA[0], avgX + 1, 6);
    WriteInDOCX ('%.2f' % weightedGPA[1], avgX + 1, 8);

    cumGPA = (float (tb.cell (avgX, 2).text) + float (tb.cell (avgX, 4).text) + float (tb.cell (avgX + 1, 6).text)) / 3;  # get averages of 9th, 10th, and 11th grades
    WriteInDOCX ('%.2f' % cumGPA, avgX + 2, 1);

    if (TrySave ()):
        print ("");
        print ("");
        print ("");
        print ("--------------------------------------------------------------");
        print ("-------------------- Conversion Complete! --------------------");
        print ("--------------------------------------------------------------");

        showinfo (title = "Complete", message = "Conversion Complete. Please check the file, especially the total class hours and the average grades.")
        showinfo (title = "Complete", message = "Other that class hours, grades, and averages (which you should obviously check) also check your name grade etc.")
        showinfo (title = "Complete", message = "Just check everything and NOT JUST PRINT/SEND THIS. This is a dumb program and it can make mistakes!")
        showinfo (title = "Complete", message = "Also please report any mistakes to me so that I can fix them :) -Atahan")
        to2 = threading.Thread (target = OpenFinalFile)
        to2.start ()
    else:
        print ("");
        print ("");
        print ("");
        print ("--------------------------------------------------------------");
        print ("--------------------- Conversion FAILED! ---------------------");
        print ("--------------------------------------------------------------");

        showwarning (title = "Error", message = "Conversion Failed")

    buttonText.set ("Start Converting");
    button.config (state = ACTIVE);

    root.destroy();
    root.quit();
    sys.exit();


def TrySave ():
    global output;
    output = asksaveasfilename (defaultextension = ".docx");

    try:
        doc.save (output);
        return True;
    except Exception as e:
        return FileReadorWriteError (TrySave);

def OpenFinalFile ():
    os.startfile ( output )

Label (mainframe).grid (row = 19, column = 0, sticky = (EW), ipady = 10);  # spacer
button = Button (mainframe, text = "Start Converting", command = ok);
button.grid (row = 20, column = 0, sticky = (EW), ipady = 2);

showinfo (title = "Info", message = "To use this tool, you need to get a Turkish Excel Transcript file from the school administration. Sadly you can't get the file yourself from E-Okul. You need the normal Excel file, Word Transcript and 'data only' Excel won't work. Also make sure that it's Turkish.")

mainloop ()
