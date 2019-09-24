from xlrd import open_workbook
from docx import Document

input = "input.xls"
output = "output.docx"

wb = open_workbook (input);
ws = wb.sheet_by_name("Sheet1")

x = 0;
y = 0;
for myRow in ws.get_rows():
    valuess = str(x) + " - ";
    y = 0;
    for myCell in myRow:
        valuess += "(" + str(y) + "- " + str(myCell.value) + ")";
        y+=1;

    x+=1;
    print(valuess);

print("name:" + str(ws.cell_value(6,3)));

doc = Document(output);

t1 = doc.tables[0];

paragraph = t1.cell(0,1).paragraphs[0]
paragraph.runs[0].text = str(ws.cell_value(6,3)).title();

doc.save(output);
